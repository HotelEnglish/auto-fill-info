import os
from docx import Document
import re
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client
import pythoncom

def convert_doc_to_docx(doc_path):
    """将.doc文件转换为.docx格式"""
    try:
        # 初始化 Word 应用
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(doc_path)
        docx_path = doc_path + "x"  # 添加 x 后缀
        doc.SaveAs2(docx_path, FileFormat=16)  # 16 表示 docx 格式
        doc.Close()
        word.Quit()
        return docx_path
    except Exception as e:
        print(f"转换文件 {doc_path} 时出错: {str(e)}")
        return None
    finally:
        pythoncom.CoUninitialize()

def read_personal_info(info_file):
    """读取个人信息文档"""
    if not os.path.exists(info_file):
        print(f"错误: 文件 '{info_file}' 不存在!")
        return None
        
    try:
        doc = Document(info_file)
        info_dict = {}
        
        # 读取文档中的文本
        for paragraph in doc.paragraphs:
            line = paragraph.text.strip()
            if ':' in line:
                key, value = line.split(':', 1)
                info_dict[key.strip()] = value.strip()
        
        # 添加常见的别名映射
        aliases = {
            '姓名': ['名字', '填表人', '申报人', '本人姓名'],
            '性别': ['性别'],
            '年龄': ['年龄', '岁数'],
            '出生年月': ['出生日期', '生日'],
            '工作单位': ['单位', '所在单位', '工作单位名称'],
            '职务': ['现任职务', '担任职务'],
            '职称': ['现职称', '现任职称', '现有职称'],
            '联系电话': ['电话', '手机号码', '联系方式'],
            # 添加更多别名映射...
        }
        
        # 扩展信息字典，添加别名
        expanded_dict = {}
        for key, value in info_dict.items():
            expanded_dict[key] = value
            if key in aliases:
                for alias in aliases[key]:
                    expanded_dict[alias] = value
                    
        return expanded_dict
        
    except Exception as e:
        print(f"读取个人信息文件时出错: {str(e)}")
        return None

def find_fillable_fields(text):
    """识别需要填写的字段"""
    # 常见的填写标记模式
    patterns = [
        # 表格中的标准字段
        r'(姓名)$',  # 表格中的"姓名"字段
        r'(性别)$',  # 表格中的"性别"字段
        r'(出生年月(?:\s*（[^）]*）)?)$',  # 出生年月（带可能的括号说明）
        r'(教师资格证种类及学科)$',  # 教师资格证
        r'(身份证\s*号码)$',  # 身份证号码
        r'(毕业院校)$',  # 毕业院校
        r'(学历\s*学位)$',  # 学历学位
        r'(所学专业)$',  # 所学专业
        r'(现工作单位)$',  # 现工作单位
        r'(参加工作时间)$',  # 参加工作时间
        r'(任教学科)$',  # 任教学科
        # 其他通用模式
        r'(\{[^}]+\})',  # 花括号包围的内容
        r'(\[[^\]]+\])',  # 方括号包围的内容
        r'(<[^>]+>)',    # 尖括号包围的内容
        r'([_]{2,})',    # 连续的下划线
        r'(□.*?□)',      # 复选框之间的内容
    ]
    
    fields = []
    for pattern in patterns:
        matches = re.finditer(pattern, text.strip(), re.UNICODE)
        for match in matches:
            fields.append({
                'start': match.start(),
                'end': match.end(),
                'field': match.group(),
                'context': text.strip()  # 对于表格，整个单元格内容作为上下文
            })
    return fields

def fill_document(template_file, info_dict):
    """填写单个文档"""
    try:
        # 如果是.doc文件，先转换为.docx
        if template_file.endswith('.doc'):
            docx_file = convert_doc_to_docx(os.path.abspath(template_file))
            if not docx_file:
                return
        else:
            docx_file = template_file
            
        doc = Document(docx_file)
        modified = False
        filled_fields = set()  # 记录已填写的字段
        
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    # 检查当前单元格是否包含字段名
                    fields = find_fillable_fields(cell.text)
                    if fields:
                        for field_info in fields:
                            field = field_info['field']
                            # 查找匹配的信息
                            for key, value in info_dict.items():
                                if (key.lower() in field.lower() or 
                                    field.lower() in key.lower()):
                                    # 在下一个单元格填写信息
                                    if i + 1 < len(row.cells):
                                        next_cell = row.cells[i + 1]
                                        if not next_cell.text.strip():  # 如果单元格为空
                                            next_cell.text = value
                                            modified = True
                                            filled_fields.add(key)
                                    break
        
        if modified:
            # 保存为新文件
            output_file = f"filled_{os.path.splitext(template_file)[0]}.docx"
            doc.save(output_file)
            print(f"已完成文件填写: {output_file}")
            print("已填写的字段:")
            for field in filled_fields:
                print(f"  {field}: {info_dict[field]}")
        else:
            print(f"警告: 在 {template_file} 中没有找到可填写的字段")
            print("请确保表格中包含以下字段:")
            print("  - 姓名")
            print("  - 性别")
            print("  - 出生年月")
            print("  - 教师资格证种类及学科")
            print("  - 身份证号码")
            print("  - 毕业院校")
            print("  - 学历学位")
            print("  - 所学专业")
            print("  - 现工作单位")
            print("  - 参加工作时间")
            print("  - 任教学科")
            
    except Exception as e:
        print(f"处理文件 {template_file} 时出错: {str(e)}")

def main():
    info_file = "information.docx"
    
    # 显示当前工作目录
    print(f"当前工作目录: {os.getcwd()}")
    
    # 读取个人信息
    info_dict = read_personal_info(info_file)
    if not info_dict:
        return
    
    print("读取到的个人信息:")
    for key, value in info_dict.items():
        print(f"  {key}: {value}")
    
    # 获取当前目录下所有的文档文件
    doc_files = [f for f in os.listdir() if f.endswith(('.doc', '.docx')) 
                 and f != info_file and not f.startswith('filled_')]
    
    if not doc_files:
        print("没有找到需要处理的文档文件")
        return
        
    print(f"\n找到以下需要处理的文件:")
    for i, file in enumerate(doc_files, 1):
        print(f"{i}. {file}")
    
    # 处理每个文档
    for doc_file in doc_files:
        print(f"\n正在处理文件: {doc_file}")
        fill_document(doc_file, info_dict)

if __name__ == "__main__":
    main() 
